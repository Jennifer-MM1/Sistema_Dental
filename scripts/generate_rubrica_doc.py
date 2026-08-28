import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_rubric_document(filename):
    doc = docx.Document()

    # Set page margins (0.6 in)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Styles
    navy = RGBColor(0, 51, 102)
    white = RGBColor(255, 255, 255)

    # Header Title
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_header.add_run("UNIVERSIDAD TECNOLÓGICA DE SAN JUAN DEL RÍO\n")
    run_inst.bold = True
    run_inst.font.size = Pt(14)
    run_inst.font.name = "Calibri"
    run_inst.font.color.rgb = navy

    run_sub = p_header.add_run("Unidad Académica Jalpan • Ingeniería en Desarrollo y Gestión de Software\n")
    run_sub.bold = True
    run_sub.font.size = Pt(11)
    run_sub.font.name = "Calibri"

    run_materia = p_header.add_run("DESARROLLO PARA DISPOSITIVOS INTELIGENTES — EVALUACIÓN MULTI-PERSPECTIVA\n")
    run_materia.bold = True
    run_materia.font.size = Pt(12)
    run_materia.font.color.rgb = navy

    run_doc_title = p_header.add_run("RÚBRICAS DE EVALUACIÓN DE PROYECTO FINAL (DENTALSYNC)\n")
    run_doc_title.bold = True
    run_doc_title.font.size = Pt(13)

    # Info Table
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(info_table, color="003366", sz="8")

    labels = [
        [("Alumnos:", " T.S.U. Montes Rodríguez Oliver Edson / T.S.U. Villa Aguillón José Manuel"),
         ("Materia:", " Desarrollo para Dispositivos Inteligentes (9°)")],
        [("Proyecto:", " VitalCare Link / DentalSync (Ecosistema Clínico)"),
         ("Profesor / Evaluador:", " ______________________________________")]
    ]

    for r_idx in range(2):
        for c_idx in range(2):
            cell = info_table.rows[r_idx].cells[c_idx]
            set_cell_background(cell, "F3F7FA")
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            lbl, val = labels[r_idx][c_idx]
            p = cell.paragraphs[0]
            r1 = p.add_run(lbl)
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(val)
            r2.font.size = Pt(10)

    doc.add_paragraph() # spacer

    # Instructions
    p_inst = doc.add_paragraph()
    r_inst_title = p_inst.add_run("Instrucciones de Evaluación: ")
    r_inst_title.bold = True
    r_inst_title.font.size = Pt(10.5)
    r_inst_body = p_inst.add_run(
        "Para evaluar el proyecto, se evalúa cada dispositivo de forma independiente a través de su respectiva rúbrica. Cada evaluación exige la presentación interactiva ejecutando la app en vivo dentro de su emulador/dispositivo target. Los criterios están enfocados puramente en la funcionalidad visible, usabilidad, heurísticas de diseño y comportamiento del sistema ante el evaluador. Al final de cada rúbrica se incluye el apartado de calificación correspondiente."
    )
    r_inst_body.font.size = Pt(10)

    # Function to create a standard rubric section
    def create_rubric_section(title, subtitle, criteria_data, calc_data, rubric_num):
        doc.add_paragraph()
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(f"RÚBRICA {rubric_num}: {title.upper()}")
        r_t.bold = True
        r_t.font.size = Pt(12)
        r_t.font.color.rgb = navy

        p_st = doc.add_paragraph()
        r_st = p_st.add_run(subtitle)
        r_st.font.size = Pt(9.5)
        r_st.font.italic = True

        # Main Rubric Table
        headers = ["Criterio y Enfoque", "Sobresaliente\n(100% - 90%)", "Satisfactorio\n(89% - 80%)", "Suficiente\n(79% - 70%)", "Insuficiente\n(< 70%)"]
        col_widths = [Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.3), Inches(1.2)]

        table = doc.add_table(rows=1 + len(criteria_data), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color="003366", sz="6")

        # Header row
        hdr_cells = table.rows[0].cells
        for i, head_text in enumerate(headers):
            hdr_cells[i].width = col_widths[i]
            set_cell_background(hdr_cells[i], "003366")
            set_cell_margins(hdr_cells[i], top=140, bottom=140, left=100, right=100)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(head_text)
            run.bold = True
            run.font.color.rgb = white
            run.font.size = Pt(9.5)

        # Data rows
        for row_idx, row_data in enumerate(criteria_data):
            row_cells = table.rows[row_idx + 1].cells
            bg_color = "FFFFFF" if row_idx % 2 == 0 else "F9FBFD"
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].width = col_widths[col_idx]
                set_cell_background(row_cells[col_idx], bg_color)
                set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=100, right=100)
                p = row_cells[col_idx].paragraphs[0]
                if col_idx == 0:
                    run = p.add_run(cell_value[0] + "\n")
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = navy
                    run2 = p.add_run(cell_value[1])
                    run2.font.size = Pt(8.5)
                    run2.font.italic = True
                else:
                    run = p.add_run(cell_value)
                    run.font.size = Pt(8.5)

        # Score Table
        doc.add_paragraph()
        p_sc_hdr = doc.add_paragraph()
        r_sc_hdr = p_sc_hdr.add_run(f"APARTADO DE CALIFICACIÓN — RÚBRICA {rubric_num} ({title.upper()})")
        r_sc_hdr.bold = True
        r_sc_hdr.font.size = Pt(10.5)
        r_sc_hdr.font.color.rgb = navy

        sc_table = doc.add_table(rows=1 + len(calc_data) + 1 + 1, cols=4)
        sc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(sc_table, color="003366", sz="6")

        sc_headers = ["Criterio Evaluado", "Ponderación", "Puntaje Obtenido (0-100)", "Calificación Ponderada"]
        sc_hdr_cells = sc_table.rows[0].cells
        for i, text in enumerate(sc_headers):
            set_cell_background(sc_hdr_cells[i], "006C9C")
            set_cell_margins(sc_hdr_cells[i], top=120, bottom=120, left=100, right=100)
            p = sc_hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = white
            run.font.size = Pt(9)

        for idx, (crit, pond) in enumerate(calc_data):
            row_cells = sc_table.rows[idx + 1].cells
            set_cell_margins(row_cells[0], top=80, bottom=80, left=100, right=100)
            p = row_cells[0].paragraphs[0]
            r = p.add_run(crit)
            r.font.size = Pt(9)

            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_pond = row_cells[1].paragraphs[0].add_run(pond)
            r_pond.font.size = Pt(9)

        # Total Row
        tot_row = sc_table.rows[len(calc_data) + 1].cells
        set_cell_background(tot_row[0], "E6F0F8")
        set_cell_background(tot_row[1], "E6F0F8")
        set_cell_background(tot_row[2], "E6F0F8")
        set_cell_background(tot_row[3], "E6F0F8")

        p_tot = tot_row[0].paragraphs[0]
        r_tot = p_tot.add_run("CALIFICACIÓN TOTAL DISPOSITIVO")
        r_tot.bold = True
        r_tot.font.size = Pt(9.5)

        tot_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_tot_p = tot_row[1].paragraphs[0].add_run("100%")
        r_tot_p.bold = True

        tot_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_prom = tot_row[2].paragraphs[0].add_run("Promedio Final:")
        r_prom.bold = True

        tot_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_score = tot_row[3].paragraphs[0].add_run("     / 100")
        r_score.bold = True

        # Retro row
        obs_row = sc_table.rows[len(calc_data) + 2].cells
        obs_cell = obs_row[0]
        for c in range(1, 4):
            obs_cell.merge(obs_row[c])
        set_cell_margins(obs_cell, top=100, bottom=300, left=100, right=100)
        p_obs = obs_cell.paragraphs[0]
        r_obs = p_obs.add_run("Observaciones y Retroalimentación del Dispositivo:\n\n")
        r_obs.bold = True
        r_obs.font.size = Pt(9)

    # -------------------------------------------------------------
    # DATA 1: SMARTWATCH (WEAR OS)
    # -------------------------------------------------------------
    smartwatch_criteria = [
        (
            ("Reconocimiento y Simplicidad de Contexto", "Heurísticas"),
            "Muestra exclusivamente información crítica (Código de Turno y Consultorio) de inmediato. Evita exigir memoria de reconocimiento o navegación compleja en pantalla pequeña.",
            "Muestra el turno, pero agrega textos secundarios largos que requieren hacer scroll en la pantalla del reloj inteligente.",
            "Falta de contexto claro (ej. solo muestra '¡Es tu turno!' sin especificar consultorio ni código de turno para verificar).",
            "La interfaz del reloj se congela al intentar leer datos o no despliega el aviso de turno."
        ),
        (
            ("Accesibilidad Sensorial y Multimodal", "Accesibilidad"),
            "Uso redundante de canales: color verde de alta visibilidad para llamada, icono gigante de campana y alerta hápica (vibración del reloj) para asegurar la percepción del paciente.",
            "Tiene vibración, pero la pantalla no aprovecha el color de contraste o el tamaño del icono es pequeño para ver de reojo.",
            "Sin alertas vibratorias, con texto pequeño sobre fondos de bajo contraste que pasan totalmente desapercibidos.",
            "Alertas ilegibles, sin respuesta háptica o desproporcionadas para la pantalla circular del reloj."
        ),
        (
            ("Coherencia y Sincronización Inmediata", "Coherencia"),
            "Transiciona instantáneamente de estado en perfecta sincronía con la tablet del dentista. Mantiene consistencia estricta en el nombre del consultorio y código de turno.",
            "La alerta se recibe con retraso (más de 10 segundos) o la nomenclatura varía ligeramente de los demás canales de comunicación.",
            "Falta de sincronización, mostrando información desactualizada o de una cita previa debido a problemas de caché.",
            "No reacciona al turno actual ni sincroniza eventos con el sistema."
        ),
        (
            ("Eficacia y Ergonomía del Botón de Confirmación", "UX y UI"),
            "El botón 'CONFIRMAR / EN CAMINO' abarca un área de toque generosa en la base de la pantalla, evitando pulsaciones accidentales mientras se camina al consultorio.",
            "El botón está disponible, pero es pequeño o está muy cerca de los bordes del dispositivo, dificultando el toque preciso.",
            "No incluye botón de confirmación en el reloj (forzando a sacar el teléfono) o el botón es tan pequeño que es casi imposible de presionar.",
            "Botones diminutos e imposibles de presionar en el reloj. Texto ilegible o desbordado fuera del área circular de la pantalla."
        )
    ]
    smartwatch_score = [
        ("RECONOCIMIENTO Y SIMPLICIDAD DE CONTEXTO (HEURÍSTICAS)", "30%"),
        ("ACCESIBILIDAD SENSORIAL Y MULTIMODAL (ACCESIBILIDAD)", "25%"),
        ("COHERENCIA Y SINCRONIZACIÓN INMEDIATA (COHERENCIA)", "25%"),
        ("EFICACIA Y ERGONOMÍA DEL BOTÓN DE CONFIRMACIÓN (UX/UI)", "20%")
    ]
    create_rubric_section(
        "Smartwatch (Wear OS)",
        "Perspectiva: Especialización Técnica | Ponderación: Enfoque Técnico (50%) • Estética/UX (25%) • Impacto/Gestión (25%)\nEvalúa la micro-interfaz en el reloj inteligente del paciente orientada a notificaciones hápicas rápidas cuando es su turno de ingresar al consultorio.",
        smartwatch_criteria,
        smartwatch_score,
        1
    )

    # -------------------------------------------------------------
    # DATA 2: SMARTPHONE (ANDROID MÓVIL)
    # -------------------------------------------------------------
    smartphone_criteria = [
        (
            ("Diseño & UX: Visualización de Datos y Lectura Accesible", "Diseño & UX (50%)"),
            "Diseño moderno y estructurado. Tarjetas de servicios y citas muy claras e intuitivas. Integración de descripciones verbales sintéticas activas para lectores de pantalla.",
            "Diseño visual estético y organizado mediante tarjetas de información. Navegación fluida entre pantallas y soporte básico para lectura de texto por voz.",
            "Vista de citas difícil de interpretar o saturada. El lector de pantalla pronuncia datos numéricos aislados sin explicar el contexto.",
            "Diseño desalineado, tipografía ilegible o colores sin jerarquía. Inaccesible o confuso para el usuario."
        ),
        (
            ("Técnico: Sincronización en Vivo y Notificaciones Push", "Técnico (25%)"),
            "Indicador de conectividad en tiempo real siempre visible ('Sincronizado hace 2 min'). Notificación push inmediata en pantalla ante llamado de turno o avisos de la clínica.",
            "Sincronización correcta de datos e historial de citas. Módulo de avisos responde bien al actualizar o consultar eventos pasados.",
            "Actualización de datos con demoras perceptibles en la pantalla. Notificaciones intermitentes ante cambios de estado.",
            "Fallas de conexión continuas. No despliega el historial de citas ni recibe notificaciones durante la prueba."
        ),
        (
            ("Gestión / Impacto: Prevención de Errores y Registro QR", "Gestión / Impacto (25%)"),
            "El flujo de check-in con código QR y geolocalización es intuitivo y rápido. Previene reservas duplicadas deshabilitando visualmente horarios pasados o ya ocupados.",
            "El autoregistro por QR funciona, pero requiere demasiados pasos o no da retroalimentación clara de confirmación en pantalla.",
            "El proceso de check-in es confuso, propenso a fallar sin explicación y genera fricción al llegar a la clínica.",
            "Permite cometer errores de agendamiento (citas empalmadas) o no informa el estado actual del paciente en sala de espera."
        ),
        (
            ("Claridad de Llamadas a la Acción (CTA)", "UX y UI"),
            "Los botones clave ('Agendar Cita', 'Escanear QR', 'Ver Historial') son visualmente prominentes, fáciles de pulsar y están ubicados de manera óptima para el alcance del pulgar.",
            "Los botones de acción son identificables, pero su ubicación no es la más ergonómica o compiten visualmente con elementos secundarios.",
            "Los botones de acción principal son difíciles de encontrar, tienen tamaños de toque menores a 48x48 dp o causan pulsaciones accidentales.",
            "Botones diminutos o ausentes que impiden completar la acción principal."
        )
    ]
    smartphone_score = [
        ("DISEÑO & UX: VISUALIZACIÓN Y LECTURA ACCESIBLE", "50%"),
        ("TÉCNICO: SINCRONIZACIÓN EN VIVO Y NOTIFICACIONES", "25%"),
        ("GESTIÓN / IMPACTO: PREVENCIÓN DE ERRORES Y REGISTRO QR", "25%")
    ]
    create_rubric_section(
        "Smartphone (Android Móvil)",
        "Perspectiva: Creativa y Experiencia de Usuario | Ponderación: Enfoque Técnico (25%) • Estética/UX (50%) • Impacto/Gestión (25%)\nEvalúa la experiencia del paciente en su dispositivo móvil al agendar citas y realizar el check-in presencial en la clínica dental.",
        smartphone_criteria,
        smartphone_score,
        2
    )

    # -------------------------------------------------------------
    # DATA 3: TABLET / ESCRITORIO (DENTISTA / ADMINISTRADOR)
    # -------------------------------------------------------------
    dashboard_criteria = [
        (
            ("Control, Libertad y Sincronización Realtime", "Heurísticas"),
            "El usuario tiene control total sobre el orden de la fila (Llamar, En Consulta, Completar, Cancelar). La actualización es instantánea mediante Supabase Realtime con botón de recarga 🔄.",
            "La fila se actualiza automáticamente, pero las acciones de control ('Llamar', 'Completar') están ocultas en menús o no permiten reversión rápida.",
            "La tabla requiere actualización manual o no permite modificar el estado del paciente, forzando cierres o recargas de pantalla.",
            "No permite modificar el estado del paciente o causa bloqueos al gestionar la atención del día."
        ),
        (
            ("Diseño Responsivo, Ajuste y Desplazamiento Fluido", "Accesibilidad / UX"),
            "Diseño responsivo adaptado a pantalla de tablet y escritorio. Calendario Mensual/Semanal completo y Directorio de Pacientes con desplazamiento vertical y horizontal 100% fluido sin recortes.",
            "Diseñado según la Ley de Fitts: botones de acción sobredimensionados para pulsación rápida, pero requiere desplazarse en pantallas de menor tamaño.",
            "Botones pequeños y muy juntos, propensos a toques accidentales y difíciles de presionar en situaciones de ritmo rápido.",
            "Diseño caótico, textos desalineados, tablas desbordadas y terminología desordenada que dificulta la operación."
        ),
        (
            ("Coherencia y Estructura Organizativa en Panel Clínico", "Coherencia"),
            "Organización estructurada del layout por módulos (Dashboard, Citas Activas, Pacientes Clínicos, Historial Clínico). Los códigos de color de estado coinciden con todo el sistema.",
            "Usa los mismos colores, pero la distribución de la tabla es confusa en pantallas de menor tamaño o cambia de ubicación los elementos de navegación.",
            "Inconsistencia visual entre módulos o superposición de textos en la lista de atención.",
            "Ausencia de jerarquía visual; mezcla notas clínicas con la cola de recepción de manera desordenada."
        ),
        (
            ("Eficiencia y Reducción de Carga Cognitiva", "UX y UI"),
            "Muestra información crítica (Nombre, Servicio, Estado) resumida y legible de un vistazo. Minimiza el desorden visual para evitar la fatiga bajo las luces clínicas. Acceso rápido a Odontograma y PDF.",
            "Muestra la información de la fila, pero la densidad es excesiva, requiriendo esfuerzo mental para identificar al siguiente paciente.",
            "Sobrecarga de información irrelevante, fuentes muy pequeñas y falta de jerarquía visual que satura al personal.",
            "Satura al personal clínico con información irrelevante e impide acceder fácilmente a la historia clínica."
        ),
        (
            ("Diseño de Dashboard y Visualización de Métricas", "UX y UI"),
            "Métricas clave ('Citas Hoy', 'Completadas', 'Clientes', 'Dentistas Activos') presentadas mediante tarjetas de resumen limpias con iconos descriptivos e indicadores claros.",
            "Las métricas están presentes, pero son difíciles de leer rápidamente o no tienen una separación visual efectiva del listado de la fila.",
            "Ausencia de resumen de métricas, o representadas mediante gráficos confusos y mal escalados que no aportan valor operativo.",
            "Ausencia total de resumen de métricas o panel desconfigurado."
        )
    ]
    dashboard_score = [
        ("CONTROL, LIBERTAD Y SINCRONIZACIÓN REALTIME (HEURÍSTICAS)", "25%"),
        ("DISEÑO RESPONSIVO Y DESPLAZAMIENTO FLUIDO (ACCESIBILIDAD/UX)", "25%"),
        ("COHERENCIA Y ESTRUCTURA ORGANIZATIVA (COHERENCIA)", "20%"),
        ("EFICIENCIA Y REDUCCIÓN DE CARGA COGNITIVA (UX/UI)", "15%"),
        ("DISEÑO DE DASHBOARD Y VISUALIZACIÓN DE MÉTRICAS (UX/UI)", "15%")
    ]
    create_rubric_section(
        "Dashboard en Tablet / Escritorio (Dentista / Administrador)",
        "Perspectiva: Gestión, Producto e Impacto Social | Ponderación: Enfoque Técnico (25%) • Estética/UX (25%) • Impacto/Gestión (50%)\nEvalúa la consola de control clínico donde el dentista o recepcionista gestiona la cola de pacientes, el calendario mensual/semanal responsivo, el directorio de pacientes y el historial clínico.",
        dashboard_criteria,
        dashboard_score,
        3
    )

    # -------------------------------------------------------------
    # FINAL SUMMARY EVALUATION
    # -------------------------------------------------------------
    doc.add_paragraph()
    p_fin_hdr = doc.add_paragraph()
    r_fin_hdr = p_fin_hdr.add_run("EVALUACIÓN FINAL DEL ECOSISTEMA VITALCARE LINK / DENTALSYNC")
    r_fin_hdr.bold = True
    r_fin_hdr.font.size = Pt(13)
    r_fin_hdr.font.color.rgb = navy

    summary_table = doc.add_table(rows=5, cols=3)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(summary_table, color="003366", sz="8")

    sum_headers = ["Dispositivo Evaluado", "Ponderación Ecosistema", "Calificación Final Obtenida"]
    for i, h_text in enumerate(sum_headers):
        cell = summary_table.rows[0].cells[i]
        set_cell_background(cell, "003366")
        set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h_text)
        r.bold = True
        r.font.color.rgb = white
        r.font.size = Pt(10)

    rows_summary_data = [
        ("Rúbrica 1: Smartwatch (Wear OS)", "33.3%"),
        ("Rúbrica 2: Smartphone (Android Móvil)", "33.3%"),
        ("Rúbrica 3: Android TV / Tablet / Escritorio", "33.4%")
    ]

    for idx, (comp, pond) in enumerate(rows_summary_data):
        row_cells = summary_table.rows[idx + 1].cells
        set_cell_margins(row_cells[0], top=100, bottom=100, left=120, right=120)
        set_cell_margins(row_cells[1], top=100, bottom=100, left=120, right=120)
        set_cell_margins(row_cells[2], top=100, bottom=100, left=120, right=120)

        p0 = row_cells[0].paragraphs[0]
        r0 = p0.add_run(comp)
        r0.font.size = Pt(9.5)

        p1 = row_cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(pond)
        r1.font.size = Pt(9.5)

    # Final Total row
    tot_cells = summary_table.rows[4].cells
    set_cell_background(tot_cells[0], "006C9C")
    set_cell_background(tot_cells[1], "006C9C")
    set_cell_background(tot_cells[2], "006C9C")
    for c in tot_cells:
        set_cell_margins(c, top=140, bottom=140, left=120, right=120)

    p_f_tot = tot_cells[0].paragraphs[0]
    r_f_tot = p_f_tot.add_run("CALIFICACIÓN FINAL DEL PROYECTO")
    r_f_tot.bold = True
    r_f_tot.font.color.rgb = white
    r_f_tot.font.size = Pt(10.5)

    p_f_p = tot_cells[1].paragraphs[0]
    p_f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f_p = p_f_p.add_run("100%")
    r_f_p.bold = True
    r_f_p.font.color.rgb = white

    p_f_sc = tot_cells[2].paragraphs[0]
    p_f_sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f_sc = p_f_sc.add_run("          / 100")
    r_f_sc.bold = True
    r_f_sc.font.color.rgb = white

    # Signature Block
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_line = p_sig.add_run("_____________________________________________________\n")
    r_line.bold = True
    r_sig_lbl = p_sig.add_run("Firma del Profesor / Evaluador")
    r_sig_lbl.bold = True
    r_sig_lbl.font.size = Pt(11)
    r_sig_lbl.font.color.rgb = navy

    doc.save(filename)
    print(f"Document saved successfully as '{filename}'")

if __name__ == "__main__":
    out_path = "c:/Users/julia/Documents/proyectos/Sistema_Dental/Rubrica_Evaluacion_DentalSync_Final.docx"
    build_rubric_document(out_path)
    try:
        build_rubric_document("c:/Users/julia/Documents/proyectos/Sistema_Dental/Rubrica_Evaluacion_DentalSync.docx")
    except Exception as e:
        print("Note on primary file:", e)
